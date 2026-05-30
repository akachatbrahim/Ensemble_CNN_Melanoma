"""
Extract metrics from organized PKL files (save folder structure) and generate Word document.
Reads from save/models_lr{lr}_wd{wd}/ directory structure with organized model files.
"""

import os
import pickle
import sys
import warnings
import matplotlib.pyplot as plt
import numpy as np
import io
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sklearn.metrics import confusion_matrix
import seaborn as sns
from src.modules.memory_optimized_test import test_memory_optimized
#optimized configuration for memory-constrained systems
test_memory_optimized()


# Suppress warnings
warnings.filterwarnings('ignore')
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'

# Import config

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
try:
    from src.config.Config import learning_rate, weight_decay
except ImportError:
    print("✗ Error importing Config. Verif Existing the file.")
    exit()
    #learning_rate = 1e-3
    #weight_decay = 1e-3

plt.ioff()  # Disable interactive mode


def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)


def extract_pkl_data(pkl_file):
    """Extract data from PKL file"""
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            with open(pkl_file, 'rb') as f:
                data = pickle.load(f)
        
        result = {
            'name': data.get('name', 'Unknown'),
            'metrics': {},
            'history': {},
            'y_true': None,
            'y_pred': None,
            'success': True
        }
        
        # Try to extract metrics from different possible structures
        exp_data = data.get('experiment_data', {})
        
        # First try: nested by model name
        if isinstance(exp_data, dict):
            metrics = exp_data.get('metrics', {})
            if isinstance(metrics, dict):
                # Check if metrics are nested by model name or flat
                for key, value in metrics.items():
                    if isinstance(value, dict):
                        # Nested structure: metrics[model_name][metric_key]
                        if 'y_true' in value:
                            result['y_true'] = value.get('y_true')
                        if 'y_pred' in value:
                            result['y_pred'] = value.get('y_pred')
                        
                        # Extract serializable metrics
                        for k, v in value.items():
                            if k not in ['y_true', 'y_pred']:
                                try:
                                    if isinstance(v, (int, float, np.integer, np.floating)):
                                        result['metrics'][k] = float(v)
                                    elif isinstance(v, (str, bool)):
                                        result['metrics'][k] = v
                                except:
                                    pass
                        break
                    else:
                        # Flat structure: metrics[metric_key] = value
                        try:
                            if isinstance(value, (int, float, np.integer, np.floating)):
                                result['metrics'][key] = float(value)
                            elif isinstance(value, (str, bool)):
                                result['metrics'][key] = value
                        except:
                            pass
            
            history = exp_data.get('history', {})
            if history is not None:
                # Handle Keras History object
                if hasattr(history, 'history'):
                    result['history'] = history.history
                elif isinstance(history, dict):
                    result['history'] = history
        
        return result
    except Exception as e:
        print(f"  ⚠ Error loading {pkl_file}: {e}")
        return {'success': False, 'error': str(e)}


def create_training_history_plot(history, filename):
    """Create training history plot from history data"""
    if not history:
        return None
    
    try:
        fig, axes = plt.subplots(1, 2, figsize=(12, 4))
        fig.suptitle(f'Training History - {filename}', fontsize=14, fontweight='bold')
        
        # Convert history values to lists if they are numpy arrays
        hist_data = {}
        for key, values in history.items():
            if hasattr(values, 'tolist'):
                hist_data[key] = values.tolist()
            else:
                hist_data[key] = values
        
        # Plot 1: Loss
        has_loss = False
        if 'loss' in hist_data:
            epochs = range(1, len(hist_data['loss']) + 1)
            axes[0].plot(epochs, hist_data['loss'], 'b-', label='Train Loss', linewidth=2)
            has_loss = True
            if 'val_loss' in hist_data:
                axes[0].plot(epochs, hist_data['val_loss'], 'r-', label='Val Loss', linewidth=2)
        
        if has_loss:
            axes[0].set_xlabel('Epoch')
            axes[0].set_ylabel('Loss')
            axes[0].set_title('Loss')
            axes[0].legend()
            axes[0].grid(True, alpha=0.3)
        
        # Plot 2: Accuracy
        has_acc = False
        if 'accuracy' in hist_data:
            epochs = range(1, len(hist_data['accuracy']) + 1)
            axes[1].plot(epochs, hist_data['accuracy'], 'b-', label='Train Accuracy', linewidth=2)
            has_acc = True
            if 'val_accuracy' in hist_data:
                axes[1].plot(epochs, hist_data['val_accuracy'], 'r-', label='Val Accuracy', linewidth=2)
        
        if has_acc:
            axes[1].set_xlabel('Epoch')
            axes[1].set_ylabel('Accuracy')
            axes[1].set_title('Accuracy')
            axes[1].legend()
            axes[1].grid(True, alpha=0.3)
        
        if not (has_loss or has_acc):
            return None
        
        plt.tight_layout()
        
        # Save to bytes
        img_bytes = io.BytesIO()
        plt.savefig(img_bytes, format='png', dpi=100, bbox_inches='tight')
        img_bytes.seek(0)
        plt.close(fig)
        
        return img_bytes
    except Exception as e:
        print(f"    Error creating plot: {e}")
        return None


def create_confusion_matrix_plot(y_true, y_pred, filename):
    """Create confusion matrix plot with counts and percentages"""
    if y_true is None or y_pred is None:
        return None
    
    try:
        cm = confusion_matrix(y_true, y_pred)
        
        # Calculate total and percentages
        total = cm.sum()
        percentages = cm.astype('float') / total * 100
        
        # Create annotations with both counts and percentages
        annotations = np.empty_like(cm, dtype=object)
        for i in range(cm.shape[0]):
            for j in range(cm.shape[1]):
                annotations[i, j] = f'{cm[i, j]}\n({percentages[i, j]:.2f}%)'
        
        fig, ax = plt.subplots(figsize=(8, 6))
        sns.heatmap(cm, annot=annotations, fmt='', cmap='Blues', cbar=True, ax=ax,
                    xticklabels=['Benign', 'Malignant'],
                    yticklabels=['Benign', 'Malignant'],
                    cbar_kws={'label': 'Count'})
        
        ax.set_title(f'Confusion Matrix - {filename}', fontsize=14, fontweight='bold')
        ax.set_xlabel('Predicted Label')
        ax.set_ylabel('True Label')
        
        plt.tight_layout()
        
        # Save to bytes
        img_bytes = io.BytesIO()
        plt.savefig(img_bytes, format='png', dpi=100, bbox_inches='tight')
        img_bytes.seek(0)
        plt.close(fig)
        
        return img_bytes
    except Exception as e:
        print(f"    Error creating confusion matrix: {e}")
        return None


def create_metrics_bar_plot(metrics, filename):
    """Create bar plot for metrics"""
    if not metrics:
        return None
    
    try:
        # Select main metrics for bar plot - prioritize specific metrics but use whatever is available
        metric_keys_priority = ['accuracy', 'precision', 'sensitivity', 'specificity', 'f1-score', 'auc-roc']
        
        # Get available metrics in order of priority
        available_metrics = {}
        for k in metric_keys_priority:
            if k in metrics:
                available_metrics[k] = metrics[k]
        
        # If no priority metrics found, use all metrics (excluding y_true, y_pred)
        if not available_metrics:
            available_metrics = {k: v for k, v in metrics.items() 
                               if k not in ['y_true', 'y_pred'] 
                               and isinstance(v, (int, float))}
        
        if not available_metrics:
            return None
        
        fig, ax = plt.subplots(figsize=(10, 5))
        
        metric_names = [k.replace('-', ' ').replace('auc roc', 'AUC-ROC').title() for k in available_metrics.keys()]
        metric_values = list(available_metrics.values())
        
        # Create bar plot
        num_bars = len(metric_values)
        colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f']
        bar_colors = (colors * ((num_bars // len(colors)) + 1))[:num_bars]
        
        bars = ax.bar(metric_names, metric_values, color=bar_colors)
        
        # Add value labels on bars
        for bar in bars:
            height = bar.get_height()
            # Display as percentage with 2 decimal places
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{height:.2f}%',
                   ha='center', va='bottom', fontsize=10)
        
        ax.set_title(f'Performance Metrics - {filename}', fontsize=14, fontweight='bold')
        ax.set_ylabel('Score')
        ax.set_ylim(0, max(metric_values) * 1.15 if metric_values else 1.1)
        ax.grid(axis='y', alpha=0.3)
        
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        
        # Save to bytes
        img_bytes = io.BytesIO()
        plt.savefig(img_bytes, format='png', dpi=100, bbox_inches='tight')
        img_bytes.seek(0)
        plt.close(fig)
        
        return img_bytes
    except Exception as e:
        print(f"    Error creating metrics bar plot: {e}")
        return None


def add_experiment_section(doc, filename, pkl_data):
    """Add experiment section to document"""
    
    # Experiment heading
    exp_heading = doc.add_heading(filename, level=2)
    exp_heading_run = exp_heading.runs[0]
    exp_heading_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Model name
    model_name = pkl_data.get('name', 'Unknown')
    doc.add_paragraph(f"Model: {model_name}", style='List Bullet')
    
    # Metrics table
    metrics = pkl_data.get('metrics', {})
    if metrics:
        doc.add_paragraph("Performance Metrics:", style='Heading 3')
        
        table = doc.add_table(rows=len(metrics) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = "Metric"
        header_cells[1].text = "Value"
        
        for cell in header_cells:
            set_cell_background(cell, "4472C4")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        # Data rows
        for row_idx, (metric_name, metric_value) in enumerate(sorted(metrics.items()), 1):
            row = table.rows[row_idx]
            row.cells[0].text = metric_name.replace('_', ' ').replace('-', ' ').title()
            
            if isinstance(metric_value, float):
                # Always display as percentage with 2 decimal places
                row.cells[1].text = f"{metric_value:.2f}%"
            else:
                row.cells[1].text = str(metric_value)
        
        doc.add_paragraph()
    
    # Metrics Bar Plot
    metrics_plot = create_metrics_bar_plot(metrics, filename)
    if metrics_plot:
        doc.add_paragraph("Metrics Comparison:", style='Heading 3')
        try:
            doc.add_picture(metrics_plot, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"    Error adding metrics plot: {e}")
        doc.add_paragraph()
    
    # Confusion Matrix
    y_true = pkl_data.get('y_true')
    y_pred = pkl_data.get('y_pred')
    if y_true is not None and y_pred is not None:
        cm_plot = create_confusion_matrix_plot(y_true, y_pred, filename)
        if cm_plot:
            doc.add_paragraph("Confusion Matrix:", style='Heading 3')
            try:
                doc.add_picture(cm_plot, width=Inches(4.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"    Error adding confusion matrix: {e}")
            doc.add_paragraph()
    
    # Training history plot
    history = pkl_data.get('history', {})
    if history:
        doc.add_paragraph("Training History:", style='Heading 3')
        
        plot_img = create_training_history_plot(history, filename)
        if plot_img:
            try:
                doc.add_picture(plot_img, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"    Error adding plot: {e}")
        
        doc.add_paragraph()
    
    doc.add_paragraph()  # Spacing


def create_word_document(pkl_files_dir, output_path):
    """Create Word document from all PKL files"""
    
    # Check directory
    if not os.path.exists(pkl_files_dir):
        print(f"✗ Directory not found: {pkl_files_dir}")
        return False
    
    # Get PKL files
    pkl_files = sorted([f for f in os.listdir(pkl_files_dir) if f.endswith('.pkl')])
    
    if not pkl_files:
        print(f"✗ No PKL files found in {pkl_files_dir}")
        return False
    
    print(f"\n📁 Found {len(pkl_files)} PKL files")
    print(f"  Processing files...")
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('PKL Models Extraction Report (Organized Structure)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Models & Metrics from Organized Save Directory')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.italic = True
    subtitle_format.font.size = Pt(12)
    
    doc.add_paragraph(f"Configuration: Learning Rate = {learning_rate}, Weight Decay = {weight_decay}")
    doc.add_paragraph(f"Total Experiments: {len(pkl_files)}")
    doc.add_page_break()
    
    # Process each PKL file
    for pkl_file in pkl_files:
        filename = pkl_file.replace('.pkl', '')
        filepath = os.path.join(pkl_files_dir, pkl_file)
        
        print(f"  ✓ Processing: {filename}")
        
        pkl_data = extract_pkl_data(filepath)
        if pkl_data.get('success', False):
            add_experiment_section(doc, filename, pkl_data)
        else:
            doc.add_heading(filename, level=2)
            doc.add_paragraph(f"⚠ Error: {pkl_data.get('error', 'Unknown error')}")
            doc.add_paragraph()
    
    # Save document
    try:
        doc.save(output_path)
        print(f"\n✓ Word document created: {output_path}")
        return True
    except Exception as e:
        print(f"\n✗ Error saving document: {e}")
        return False


def find_organized_models_dirs():
    """Find all organized models directories in save folder"""
    base_save_dir = Path("/home/akachat/tf_env/Ensemble_CNN_Melanoma/save")
    
    if not base_save_dir.exists():
        return []
    
    # Find all models_lr*_wd* directories
    organized_dirs = sorted(base_save_dir.glob("models_lr*_wd*"))
    return organized_dirs


def select_models_directory():
    """Allow user to select which models directory to use"""
    organized_dirs = find_organized_models_dirs()
    
    if not organized_dirs:
        print("✗ No organized models directories found in /save/")
        print("  Available directories: models_lr*_wd*")
        return None
    
    print("\n📁 Available organized models directories:\n")
    for idx, dir_path in enumerate(organized_dirs, 1):
        pkl_count = len(list(dir_path.glob("*.pkl")))
        print(f"  {idx}. {dir_path.name}")
        print(f"     └─ {pkl_count} PKL files")
    
    print(f"\n  0. Back\n")
    
    choice = input("Select directory (0-" + str(len(organized_dirs)) + "): ").strip()
    
    try:
        choice_idx = int(choice)
        if choice_idx == 0:
            return None
        elif 1 <= choice_idx <= len(organized_dirs):
            return organized_dirs[choice_idx - 1]
        else:
            print("✗ Invalid choice")
            return None
    except ValueError:
        print("✗ Invalid input")
        return None


def main():
    """Main execution function"""
    import sys
    
    print("\n" + "="*80)
    print("  📊 EXTRACT ORGANIZED PKL DATA TO WORD DOCUMENT")
    print("="*80)
    
    # Check if directory is passed as argument
    pkl_dir = None
    
    if len(sys.argv) > 1:
        # Directory passed as argument
        pkl_dir = Path(sys.argv[1])
        if not pkl_dir.exists():
            print(f"✗ Directory not found: {pkl_dir}")
            return
    else:
        # Let user select directory
        pkl_dir = select_models_directory()
        
        if pkl_dir is None:
            print("✗ No directory selected")
            return
    
    print(f"\n✓ Selected directory: {pkl_dir.name}")
    
    output_dir = Path("/home/akachat/tf_env/Ensemble_CNN_Melanoma/reports")
    
    # Create output directory
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Extract config from directory name: models_lr0.0001_wd0.001
    dir_name = pkl_dir.name
    try:
        # Parse learning_rate and weight_decay from directory name
        parts = dir_name.split('_')
        lr_part = parts[1]  # lr0.0001
        wd_part = parts[2]  # wd0.001
        
        lr_value = float(lr_part[2:])  # Remove 'lr' prefix
        wd_value = float(wd_part[2:])  # Remove 'wd' prefix
        
        lr_str = f"{lr_value:.0e}".replace('e-0', 'e-').replace('e-', 'e_neg')
        wd_str = f"{wd_value:.0e}".replace('e-0', 'e-').replace('e-', 'e_neg')
    except:
        lr_str = "unknown_lr"
        wd_str = "unknown_wd"
    
    output_filename = f"PKL_Organized_Report_{dir_name}.docx"
    output_path = output_dir / output_filename
    
    # Create document
    success = create_word_document(str(pkl_dir), str(output_path))
    
    if success:
        print(f"\n✓ Report generated successfully!")
        print(f"  Location: {output_path}")
    
    print("="*80 + "\n")


if __name__ == "__main__":
    main()
