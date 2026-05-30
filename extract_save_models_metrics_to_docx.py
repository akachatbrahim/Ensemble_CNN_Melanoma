#!/usr/bin/env python3

import os
import sys
import pickle
from pathlib import Path
from datetime import datetime
from collections import defaultdict

from src.config.Config import learning_rate, weight_decay, path_project

# Project root
PROJECT_ROOT = Path(path_project).resolve()

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Installing python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx -q")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT


def shade_cell(cell, color):
    """Shade a table cell with RGB color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement('w:{}'.format(edge))
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    
    tcPr.append(tcBorders)


# ============================================================
# SELECT MODEL FROM /save
# ============================================================

def list_experiments(save_dir):
    """List all experiment folders inside /save."""
    
    save_dir = Path(save_dir)
    folders = [f for f in save_dir.iterdir() if f.is_dir()]
    folders = sorted(folders)
    
    if not folders:
        print("❌ No experiment folders found in /save")
        return []
    
    print("\n📂 Available Experiments:\n")
    for i, folder in enumerate(folders, 1):
        print(f"  [{i}] {folder.name}")
    
    return folders

def select_experiment(folders):
    print(f"\n  [{len(folders)+1}] ALL experiments")
    print("  [0] 🔙 Back to main menu")
    
    while True:
        try:
            choice = int(input("\n👉 Select experiment: "))
            
            if choice == len(folders)+1:
                print("\n✅ Selected: ALL experiments")
                return folders
            
            elif 1 <= choice <= len(folders):
                selected = folders[choice - 1]
                print(f"\n✅ Selected: {selected.name}")
                return [selected]
            elif choice == 0:
                print("  0. 🔙 Back to main menu\n")
                return None   
            else:
                print("❌ Invalid choice.")
        
        except ValueError:
            print("❌ Enter a number.")


# ============================================================
# METRICS EXTRACTION
# ============================================================

class PickleMetricsExtractor:

    def __init__(self):
        self.models_data = {}
        self.errors = []

    def extract_from_pkl(self, path):
        try:
            with open(path, "rb") as f:
                data = pickle.load(f)
            
            if isinstance(data, dict):
                if "experiment_data" in data:
                    data = data["experiment_data"]

                metrics = data.get("metrics", {})

                # ✅ cas imbriqué (CNN, ResNet, etc.)
                if isinstance(metrics, dict) and len(metrics) == 1:
                    metrics = list(metrics.values())[0]
                
                clean = {}
                for k, v in metrics.items():
                    if isinstance(v, (int, float)):
                        clean[k] = v
                
                return clean

        except Exception as e:
            self.errors.append(f"{path.name}: {str(e)}")

        return {}

    def extract_all_models(self, folder):

        model_dirs = [d for d in Path(folder).iterdir() if d.is_dir()]

        print(f"\n📂 Found {len(model_dirs)} model folders")

        for model_dir in model_dirs:

            pkl_files = list(model_dir.rglob("*.pkl"))

            if not pkl_files:
                print(f"⚠️ No PKL in {model_dir.name}")
                continue

            pkl_path = pkl_files[0]  # ✅ 1 seul fichier attendu

            model_name = model_dir.name.split("_", 1)[-1]  # ✅ nom propre

            print(f"▪ {model_name}...", end=" ")

            metrics = self.extract_from_pkl(pkl_path)

            if metrics:
                self.models_data[model_name] = metrics
                print("✓")
            else:
                print("⊘")

        return self.models_data


# ============================================================
# WORD REPORT
# ============================================================

def create_word_report(models_data, errors, output_path):

    """Create a comprehensive Word document with all metrics."""
    doc = Document()
    
    # ========================================================================
    # TITLE AND METADATA
    # ========================================================================
    title = doc.add_heading('Melanoma Detection Project', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Models - Performance Metrics Report', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.autofit = False
    meta_table.allow_autofit = False
    
    meta_rows = meta_table.rows
    
    # Generated
    meta_rows[0].cells[0].text = "Generated"
    meta_rows[0].cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    # Project
    meta_rows[1].cells[0].text = "Project"
    meta_rows[1].cells[1].text = "Ensemble CNN Melanoma Detection"
    
    # Source
    meta_rows[2].cells[0].text = "Source"
    meta_rows[2].cells[1].text = "Models folder (pickle files)"
    
    # Total models
    meta_rows[3].cells[0].text = "Total Models Processed"
    meta_rows[3].cells[1].text = str(len(models_data))
    
    # Format metadata table
    for row in meta_rows:
        shade_cell(row.cells[0], 'D9E1F2')
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # ========================================================================
    # SUMMARY SECTION
    # ========================================================================
    doc.add_heading('Executive Summary', level=1)
    
    summary_para = doc.add_paragraph()
    summary_para.add_run(f"This report contains performance metrics for ").font.size = Pt(11)
    summary_para.add_run(f"{len(models_data)}").bold = True
    summary_para.add_run(f" machine learning models trained for melanoma detection.").font.size = Pt(11)
    
    if models_data:
        # Calculate best models for key metrics
        best_accuracy = max([(name, metrics.get('accuracy', 0)) 
                            for name, metrics in models_data.items()], 
                           key=lambda x: x[1])
        
        best_auc = max([(name, metrics.get('auc-roc', 0)) 
                       for name, metrics in models_data.items()], 
                      key=lambda x: x[1])
        
        summary_list = doc.add_paragraph(style='List Bullet')
        summary_list.add_run(f"Best Accuracy: ").bold = True
        acc_val = best_accuracy[1]
        acc_str = f"{acc_val * 100:.2f}%" if acc_val <= 1.0 else f"{acc_val:.2f}%"
        summary_list.add_run(f"{best_accuracy[0]} ({acc_str})")
        
        summary_list = doc.add_paragraph(style='List Bullet')
        summary_list.add_run(f"Best AUC-ROC: ").bold = True
        auc_val = best_auc[1]
        auc_str = f"{auc_val * 100:.2f}%" if auc_val <= 1.0 else f"{auc_val:.2f}%"
        summary_list.add_run(f"{best_auc[0]} ({auc_str})")
        
        summary_list = doc.add_paragraph(style='List Bullet')
        summary_list.add_run(f"Models Processed: ").bold = True
        summary_list.add_run(f"{len(models_data)}")
        
        if errors:
            summary_list = doc.add_paragraph(style='List Bullet')
            summary_list.add_run(f"Models with errors: ").bold = True
            summary_list.add_run(f"{len(errors)}")
    
    doc.add_paragraph()  # Spacing
    
    # ========================================================================
    # DETAILED METRICS TABLE
    # ========================================================================
    doc.add_heading('Detailed Model Metrics', level=1)
    
    if not models_data:
        doc.add_paragraph("❌ No metrics data could be extracted.")
        doc.save(output_path)
        return
    
    # Determine all metric columns with specific order
    all_metrics = set()
    for metrics in models_data.values():
        all_metrics.update(metrics.keys())
    
    # Define metric order
    metric_order = ['accuracy', 'sensitivity', 'specificity', 'precision', 'f1_score', 'auc_roc', 'auc']
    metric_cols = [m for m in metric_order if m in all_metrics]
    # Add any remaining metrics not in the predefined order
    metric_cols.extend([m for m in sorted(all_metrics) if m not in metric_cols])
    
    # Create main metrics table
    num_cols = 1 + len(metric_cols)  # Model name + metrics
    table = doc.add_table(rows=len(models_data) + 1, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Model Name"
    
    for idx, metric in enumerate(metric_cols, 1):
        header_cells[idx].text = metric.upper()
    
    # Format header
    for cell in header_cells:
        shade_cell(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    for row_idx, (model_name, metrics) in enumerate(sorted(models_data.items()), 1):
        row = table.rows[row_idx]
        
        # Model name
        row.cells[0].text = model_name
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Metrics
        for col_idx, metric in enumerate(metric_cols, 1):
            value = metrics.get(metric, 0)
            
            if isinstance(value, float):
                # Convert to percentage with 2 decimal places
                if value <= 1.0:
                    row.cells[col_idx].text = f"{value * 100:.2f}%"
                else:
                    row.cells[col_idx].text = f"{value:.2f}%"
            else:
                row.cells[col_idx].text = str(value)
            
            row.cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # ========================================================================
    # METRICS STATISTICS
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('Performance Statistics', level=1)
    
    if models_data and metric_cols:
        stats_table = doc.add_table(rows=len(metric_cols) + 1, cols=5)
        stats_table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = stats_table.rows[0].cells
        headers = ['Metric', 'Min', 'Max', 'Average', 'Best Model']
        for idx, header_text in enumerate(headers):
            header_cells[idx].text = header_text
            shade_cell(header_cells[idx], '4472C4')
            for run in header_cells[idx].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Calculate stats for each metric
        for m_idx, metric in enumerate(metric_cols, 1):
            values = [metrics.get(metric, 0) for metrics in models_data.values()]
            values = [v for v in values if isinstance(v, (int, float))]
            
            if values:
                min_val = min(values)
                max_val = max(values)
                avg_val = sum(values) / len(values)
                
                # Find best model
                best_model = max(models_data.items(), 
                               key=lambda x: x[1].get(metric, 0))[0]
                
                row = stats_table.rows[m_idx]
                row.cells[0].text = metric.upper()
                min_str = f"{min_val * 100:.2f}%" if min_val <= 1.0 else f"{min_val:.2f}%"
                max_str = f"{max_val * 100:.2f}%" if max_val <= 1.0 else f"{max_val:.2f}%"
                avg_str = f"{avg_val * 100:.2f}%" if avg_val <= 1.0 else f"{avg_val:.2f}%"
                row.cells[1].text = min_str
                row.cells[2].text = max_str
                row.cells[3].text = avg_str
                row.cells[4].text = best_model
    
    doc.add_paragraph()  # Spacing
    
    # ========================================================================
    # ERRORS SECTION (if any)
    # ========================================================================
    if errors:
        doc.add_page_break()
        doc.add_heading('Processing Errors', level=1)
        
        error_intro = doc.add_paragraph()
        error_intro.add_run(f"The following {len(errors)} model(s) had errors during processing:")
        
        for error_msg in errors:
            doc.add_paragraph(error_msg, style='List Bullet')
    
    # ========================================================================
    # FOOTER
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('Report Information', level=1)
    
    footer_text = doc.add_paragraph()
    footer_text.add_run("Generated by: ").bold = True
    footer_text.add_run("PKL Metrics Extraction Script\n")
    
    footer_text = doc.add_paragraph()
    footer_text.add_run("Date: ").bold = True
    footer_text.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    
    footer_text = doc.add_paragraph()
    footer_text.add_run("Source: ").bold = True
    footer_text.add_run(f"{PROJECT_ROOT}/Models/\n")
    
    footer_text = doc.add_paragraph()
    footer_text.add_run("Format: ").bold = True
    footer_text.add_run("Microsoft Word 2007+ (.docx)\n")
    
    # Save document
    try:
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"❌ Error saving document: {e}")
        return False

# ============================================================
# MAIN
# ============================================================

def main():

    print("\n" + "="*60)
    print("📊 PKL METRICS EXTRACTION")
    print("="*60)

    save_dir = PROJECT_ROOT / "save"
    reports_dir = PROJECT_ROOT / "reports"
    reports_dir.mkdir(exist_ok=True)

    if not save_dir.exists():
        print("❌ /save folder not found")
        return False

    # Step 1
    experiments = list_experiments(save_dir)
    if not experiments:
        return False

    # Step 2
    selected_experiment = select_experiment(experiments)
    if selected_experiment is None:
        print("🔙 User chose to go back. Exiting main.")
        return False

    # ✅ IMPORTANT: normalize to list
    if isinstance(selected_experiment, list):
        selected_folders = selected_experiment
    else:
        selected_folders = [selected_experiment]

    extractor = PickleMetricsExtractor()
    all_data = {}

    total_pkl = 0

    # Step 3: process each folder
    for models_dir in selected_folders:

        print(f"\n📁 Processing: {models_dir}")

        if models_dir is None or not models_dir.exists():
            print(f"❌ Folder not found: {models_dir}")
            continue

        pkl_files = list(models_dir.rglob("*.pkl"))
        pkl_count = len(pkl_files)

        if pkl_count == 0:
            print("⚠️ No PKL files found")
            continue

        print(f"📊 Found {pkl_count} PKL files")
        total_pkl += pkl_count

        data = extractor.extract_all_models(models_dir)
        all_data.update(data)

    
        if not all_data:
            print("❌ No data extracted")
            return False

        # ✅ Output file
        output_file = reports_dir / f"PKL_Save_Report_{models_dir.name}.docx"
        #print(output_file)   
        #exit()
        print("\n📄 Creating report...")
        create_word_report(all_data, extractor.errors, output_file)

    print(f"\n✅ Report saved: {output_file}")
    print(f"📊 Total PKL processed: {total_pkl}")

    return True

# ============================================================
# RUN      
# ============================================================

if __name__ == "__main__":
    main()