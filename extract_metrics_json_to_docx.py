#!/usr/bin/env python3
"""
📊 EXTRACT ALL METRICS FROM JSON REPORT TO FORMATTED WORD DOCUMENT
===================================================================

This script reads the existing JSON report and creates a comprehensive,
professionally formatted Word document with all model metrics.
"""

import json
import sys
from pathlib import Path
from datetime import datetime
from src.config.Config import learning_rate, weight_decay,path_project

PROJECT_ROOT = Path(path_project).resolve()

#sys.path.insert(0, str(PROJECT_ROOT))

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Installing python-docx...")
    import os
    os.system(f"{sys.executable} -m pip install python-docx -q")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT


def shade_cell(cell, color):
    """Shade a table cell with RGB color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def extract_metrics_from_json(json_data):
    """Extract and organize metrics from JSON report."""
    metrics_dict = {}
    
    if 'experiments' not in json_data:
        return metrics_dict
    
    for exp_name, exp_data in json_data['experiments'].items():
        if isinstance(exp_data, dict) and 'metrics' in exp_data:
            metrics = exp_data['metrics']
            
            if isinstance(metrics, dict):
                # Flatten metrics structure
                for model_name, model_metrics in metrics.items():
                    if isinstance(model_metrics, dict):
                        # Filter out large arrays like y_true, y_pred
                        clean = {k: v for k, v in model_metrics.items() 
                                if k not in ['y_true', 'y_pred'] and isinstance(v, (int, float))}
                        
                        if clean:
                            exp_key = exp_name
                            if exp_key not in metrics_dict:
                                metrics_dict[exp_key] = {}
                            metrics_dict[exp_key][model_name] = clean
    
    return metrics_dict


def create_professional_word_report(metrics_dict, json_report_path, output_path):
    """Create a professional Word document with formatted metrics."""
    
    doc = Document()
    
    # ========================================================================
    # COVER PAGE
    # ========================================================================
    title = doc.add_heading('MELANOMA DETECTION PROJECT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.paragraph_format
    title_format.space_before = Pt(72)
    title_format.space_after = Pt(12)
    
    subtitle = doc.add_heading('Model Performance Metrics Report', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_after = Pt(48)
    
    # Add date
    date_para = doc.add_paragraph()
    date_para.add_run(f"Report Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para_format = date_para.paragraph_format
    date_para_format.space_after = Pt(24)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add metadata box
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = 'Light Grid Accent 1'
    
    meta_data = [
        ("Project", "Ensemble CNN Melanoma Detection"),
        ("Source", "Model Pickle Files and JSON Reports"),
        ("Data Directory", str(PROJECT_ROOT / "Models")),
        ("Report File", json_report_path.name),
        ("Total Experiments", str(len(metrics_dict)))
    ]
    
    for idx, (key, value) in enumerate(meta_data):
        row = meta_table.rows[idx]
        row.cells[0].text = key
        row.cells[1].text = str(value)
        # Format first column
        shade_cell(row.cells[0], 'D9E1F2')
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ========================================================================
    # EXECUTIVE SUMMARY
    # ========================================================================
    doc.add_heading('Executive Summary', level=1)
    
    # Count total models
    total_models = sum(len(exp_models) for exp_models in metrics_dict.values())
    
    summary_para = doc.add_paragraph(f"This report presents performance metrics for {total_models} machine learning models ")
    summary_para.add_run("trained for melanoma detection").bold = True
    summary_para.add_run(f". The models are organized across {len(metrics_dict)} experiments, each utilizing different combinations of deep learning architectures and machine learning classifiers.")
    
    # Calculate best performers
    best_models = {}
    for exp_name, models in metrics_dict.items():
        for model_name, metrics in models.items():
            for metric_name, value in metrics.items():
                if metric_name not in best_models:
                    best_models[metric_name] = (model_name, value, exp_name)
                else:
                    if value > best_models[metric_name][1]:
                        best_models[metric_name] = (model_name, value, exp_name)
    
    if best_models:
        doc.add_heading('Top Performers', level=2)
        
        # Sort by value in descending order and show top 5
        sorted_metrics = sorted(best_models.items(), 
                               key=lambda x: x[1][1], 
                               reverse=True)[:5]
        
        for metric_name, (model_name, value, exp_name) in sorted_metrics:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(f"{metric_name.upper()}: ").bold = True
            value_str = f"{value * 100:.2f}%" if value <= 1.0 else f"{value:.2f}%"
            para.add_run(f"{value_str} ")
            para.add_run(f"({model_name})").italic = True
    
    doc.add_page_break()
    
    # ========================================================================
    # DETAILED METRICS BY EXPERIMENT
    # ========================================================================
    doc.add_heading('Detailed Performance Metrics', level=1)
    
    if not metrics_dict:
        doc.add_paragraph("❌ No metrics data available in the report.")
        doc.save(output_path)
        return
    
    for exp_idx, (exp_name, models) in enumerate(sorted(metrics_dict.items()), 1):
        
        # Experiment heading
        doc.add_heading(f"{exp_idx}. {exp_name}", level=2)
        
        if not models:
            doc.add_paragraph("No models found for this experiment.")
            continue
        
        # Get all metric names for this experiment
        all_metrics = set()
        for model_metrics in models.values():
            all_metrics.update(model_metrics.keys())
        
        # Define metric order
        metric_order = ['accuracy', 'sensitivity', 'specificity', 'precision', 'f1_score', 'auc_roc', 'auc']
        metric_list = [m for m in metric_order if m in all_metrics]
        # Add any remaining metrics not in the predefined order
        metric_list.extend([m for m in sorted(all_metrics) if m not in metric_list])
        
        # Create metrics table
        num_cols = 1 + len(metric_list)
        table = doc.add_table(rows=len(models) + 1, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Model"
        
        for col_idx, metric in enumerate(metric_list, 1):
            header_cells[col_idx].text = metric.upper()
        
        # Format header
        for cell in header_cells:
            shade_cell(cell, '4472C4')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for row_idx, (model_name, metrics) in enumerate(sorted(models.items()), 1):
            row = table.rows[row_idx]
            
            row.cells[0].text = model_name
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            for col_idx, metric in enumerate(metric_list, 1):
                value = metrics.get(metric, "—")
                if isinstance(value, float):
                    # Convert to percentage with 2 decimal places
                    if value <= 1.0:
                        row.cells[col_idx].text = f"{value * 100:.2f}%"
                    else:
                        row.cells[col_idx].text = f"{value:.2f}%"
                else:
                    row.cells[col_idx].text = str(value)
                row.cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
    
    doc.add_page_break()
    
    # ========================================================================
    # COMPREHENSIVE STATISTICS
    # ========================================================================
    doc.add_heading('Comprehensive Statistics', level=1)
    
    # Aggregate all metrics across all experiments
    all_metrics_global = {}
    for exp_models in metrics_dict.values():
        for model_metrics in exp_models.values():
            for metric_name, value in model_metrics.items():
                if metric_name not in all_metrics_global:
                    all_metrics_global[metric_name] = []
                if isinstance(value, (int, float)):
                    all_metrics_global[metric_name].append(value)
    
    # Sort metrics in specific order
    metric_order = ['accuracy', 'sensitivity', 'specificity', 'precision', 'f1_score', 'auc_roc', 'auc']
    ordered_metrics = {m: all_metrics_global[m] for m in metric_order if m in all_metrics_global}
    # Add any remaining metrics not in the predefined order
    for m in sorted(all_metrics_global.keys()):
        if m not in ordered_metrics:
            ordered_metrics[m] = all_metrics_global[m]
    all_metrics_global = ordered_metrics
    
    if all_metrics_global:
        stats_table = doc.add_table(rows=len(all_metrics_global) + 1, cols=6)
        stats_table.style = 'Light Grid Accent 1'
        stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        header_cells = stats_table.rows[0].cells
        headers = ['Metric', 'Min', 'Max', 'Mean', 'Std Dev', 'Count']
        for idx, header_text in enumerate(headers):
            header_cells[idx].text = header_text
            shade_cell(header_cells[idx], '4472C4')
            for run in header_cells[idx].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            header_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Calculate statistics
        import statistics
        
        for m_idx, (metric_name, values) in enumerate(sorted(all_metrics_global.items()), 1):
            if values:
                min_val = min(values)
                max_val = max(values)
                mean_val = statistics.mean(values)
                std_val = statistics.stdev(values) if len(values) > 1 else 0
                
                row = stats_table.rows[m_idx]
                row.cells[0].text = metric_name.upper()
                min_str = f"{min_val * 100:.2f}%" if min_val <= 1.0 else f"{min_val:.2f}%"
                max_str = f"{max_val * 100:.2f}%" if max_val <= 1.0 else f"{max_val:.2f}%"
                mean_str = f"{mean_val * 100:.2f}%" if mean_val <= 1.0 else f"{mean_val:.2f}%"
                std_str = f"{std_val * 100:.2f}%" if std_val <= 1.0 else f"{std_val:.2f}%"
                row.cells[1].text = min_str
                row.cells[2].text = max_str
                row.cells[3].text = mean_str
                row.cells[4].text = std_str
                row.cells[5].text = str(len(values))
                
                # Center align
                for cell in row.cells[1:]:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ========================================================================
    # FOOTER
    # ========================================================================
    doc.add_heading('Report Information', level=1)
    
    info_para = doc.add_paragraph()
    info_para.add_run("Generated by: ").bold = True
    info_para.add_run("PKL Metrics Extraction Script (extract_models_metrics_to_docx.py)\n")
    
    info_para = doc.add_paragraph()
    info_para.add_run("Timestamp: ").bold = True
    info_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    
    info_para = doc.add_paragraph()
    info_para.add_run("Data Sources: ").bold = True
    info_para.add_run(f"JSON Report: {json_report_path.name}\n")
    info_para.add_run(f"Models Directory: Models/\n")
    
    info_para = doc.add_paragraph()
    info_para.add_run("Format: ").bold = True
    info_para.add_run("Microsoft Word 2007+ (.docx)\n")
    
    # Save document
    try:
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"❌ Error saving document: {e}")
        return False


def main():
    """Main execution."""
    print("\n" + "=" * 80)
    print("📊 EXTRACTING METRICS FROM JSON REPORT TO FORMATTED WORD")
    print("=" * 80)
    
    # Find JSON report
    json_files = list(PROJECT_ROOT.glob("report_*.json"))
    
    if not json_files:
        print("❌ No JSON report files found!")
        return False
    
    # Use most recent
    json_report_path = sorted(json_files, reverse=True)[0]
    print(f"\n📖 Using report: {json_report_path.name}")
    
    # Load JSON
    print("📖 Loading metrics from JSON...")
    try:
        with open(json_report_path, 'r') as f:
            json_data = json.load(f)
    except Exception as e:
        print(f"❌ Error loading JSON: {e}")
        return False
    
    # Extract metrics
    print("🔍 Extracting metrics...")
    metrics_dict = extract_metrics_from_json(json_data)
    
    total_models = sum(len(exp_models) for exp_models in metrics_dict.values())
    print(f"   ✓ Found {len(metrics_dict)} experiments with {total_models} models")
    
    if not metrics_dict:
        print("❌ No metrics data found in JSON file!")
        return False
    
    # Create Word document
    lr_str = f"{learning_rate:.0e}".replace('e-0', 'e-').replace('e-', 'e_neg')
    wd_str = f"{weight_decay:.0e}".replace('e-0', 'e-').replace('e-', 'e_neg')
    output_filename = f"Metrics_Report_JSON_lr{lr_str}_wd{wd_str}.docx"
    output_path = PROJECT_ROOT / "reports" / output_filename
    
    print(f"\n📄 Creating professional Word document...")
    success = create_professional_word_report(metrics_dict, json_report_path, output_path)
    
    if success:
        file_size = output_path.stat().st_size / 1024
        print(f"   ✓ Document created successfully")
        print(f"\n📊 Report Details:")
        print(f"   File: {output_path.name}")
        print(f"   Size: {file_size:.2f} KB")
        print(f"   Location: {output_path}")
        print(f"   Experiments: {len(metrics_dict)}")
        print(f"   Total Models: {total_models}")
    else:
        print("❌ Failed to create Word document")
        return False
    
    print("\n" + "=" * 80)
    print("✅ COMPLETE!")
    print("=" * 80 + "\n")
    
    return True


if __name__ == "__main__":
    try:
        success = main()
        sys.exit(0 if success else 1)
    except Exception as e:
        print(f"\n❌ FATAL ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
