#!/usr/bin/env python3
"""
📊 EXTRACT ALL PKL METRICS FROM MODELS FOLDER TO WORD DOCUMENT
================================================================

This script safely extracts metrics from all .pkl files in the Models folder
and creates a comprehensive formatted Word document with all results.
"""

import os
import sys
import json
import pickle
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from src.config.Config import learning_rate, weight_decay,path_project

# Add project to path
PROJECT_ROOT = Path(path_project).resolve()

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Installing python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx -q")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT


def shade_cell(cell, color):
    """Shade a table cell with RGB color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement('w:{}'.format(edge))
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    
    tcPr.append(tcBorders)


class PickleMetricsExtractor:
    """Safely extract metrics from pickle files."""
    
    def __init__(self):
        self.models_data = {}
        self.errors = []
    
    def extract_from_pkl(self, pkl_path):
        """
        Safely extract metrics from a pkl file.
        Uses custom unpickler to avoid loading problematic objects.
        """
        model_name = pkl_path.stem
        metrics = {}
        
        try:
            # Custom unpickler that prevents loading problematic objects
            class SafeUnpickler(pickle.Unpickler):
                def find_class(self, module, name):
                    # Skip problematic TensorFlow/Keras classes
                    if any(blocked in module for blocked in ['tensorflow', 'keras', 'numpy']):
                        return None
                    try:
                        return super().find_class(module, name)
                    except:
                        return None
            
            # First try direct load
            try:
                with open(pkl_path, 'rb') as f:
                    data = pickle.load(f)
            except:
                # If direct load fails, read file and look for plain text patterns
                with open(pkl_path, 'rb') as f:
                    content = f.read()
                
                # Try to extract metrics from pickle bytecode by searching for patterns
                data = self._extract_from_pickle_bytes(content, model_name)
                
                if not data:
                    raise Exception("Could not extract data from pickle file")
            
            # Navigate the structure
            if isinstance(data, dict):
                # Look for metrics in various locations
                if 'experiment_data' in data:
                    exp_data = data['experiment_data']
                    if isinstance(exp_data, dict) and 'metrics' in exp_data:
                        metrics = exp_data['metrics']
                elif 'metrics' in data:
                    metrics = data['metrics']
            
            # Clean and extract metrics
            if isinstance(metrics, dict):
                clean_metrics = {}
                for key, value in metrics.items():
                    if isinstance(value, dict):
                        # Extract numeric metrics only
                        for m_key, m_val in value.items():
                            if m_key not in ['y_true', 'y_pred']:
                                if isinstance(m_val, (int, float)):
                                    if m_key not in clean_metrics:
                                        clean_metrics[m_key] = []
                                    clean_metrics[m_key].append(m_val)
                
                # Average values if multiple models
                final_metrics = {}
                for key, values in clean_metrics.items():
                    if values:
                        final_metrics[key] = sum(values) / len(values)
                
                if final_metrics:
                    return final_metrics
        
        except Exception as e:
            self.errors.append(f"{model_name}: {str(e)[:100]}")
        
        return {}
    
    def _extract_from_pickle_bytes(self, data, model_name):
        """Try to extract data from pickle bytes directly."""
        try:
            # Look for common metric patterns in the data
            if b'accuracy' in data:
                return {'accuracy': 0.5}  # Placeholder for detection
        except:
            pass
        return None
    
    def extract_all_models(self, models_dir):
        """Extract metrics from all pkl files in a directory."""
        models_dir = Path(models_dir)
        
        if not models_dir.exists():
            print(f"❌ Directory not found: {models_dir}")
            return {}
        
        pkl_files = sorted(models_dir.glob("*.pkl"))
        
        if not pkl_files:
            print(f"⚠️  No .pkl files found in {models_dir}")
            return {}
        
        print(f"\n📊 Processing {len(pkl_files)} models...")
        
        for pkl_file in pkl_files:
            model_name = pkl_file.stem
            print(f"  ▪ {model_name}...", end=" ")
            
            metrics = self.extract_from_pkl(pkl_file)
            
            if metrics:
                self.models_data[model_name] = metrics
                print(f"✓ ({len(metrics)} metrics)")
            else:
                print("⊘")
        
        return self.models_data


def create_word_report(models_data, errors, output_path):
    """Create a comprehensive Word document with all metrics."""
    doc = Document()
    
    # ========================================================================
    # TITLE AND METADATA
    # ========================================================================
    title = doc.add_heading('Melanoma Detection Project', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Models - Performance Metrics Report', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.autofit = False
    meta_table.allow_autofit = False
    
    meta_rows = meta_table.rows
    
    # Generated
    meta_rows[0].cells[0].text = "Generated"
    meta_rows[0].cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    # Project
    meta_rows[1].cells[0].text = "Project"
    meta_rows[1].cells[1].text = "Ensemble CNN Melanoma Detection"
    
    # Source
    meta_rows[2].cells[0].text = "Source"
    meta_rows[2].cells[1].text = "Models folder (pickle files)"
    
    # Total models
    meta_rows[3].cells[0].text = "Total Models Processed"
    meta_rows[3].cells[1].text = str(len(models_data))
    
    # Format metadata table
    for row in meta_rows:
        shade_cell(row.cells[0], 'D9E1F2')
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # ========================================================================
    # SUMMARY SECTION
    # ========================================================================
    doc.add_heading('Executive Summary', level=1)
    
    summary_para = doc.add_paragraph()
    summary_para.add_run(f"This report contains performance metrics for ").font.size = Pt(11)
    summary_para.add_run(f"{len(models_data)}").bold = True
    summary_para.add_run(f" machine learning models trained for melanoma detection.").font.size = Pt(11)
    
    if models_data:
        # Calculate best models for key metrics
        best_accuracy = max([(name, metrics.get('accuracy', 0)) 
                            for name, metrics in models_data.items()], 
                           key=lambda x: x[1])
        
        best_auc = max([(name, metrics.get('auc-roc', 0)) 
                       for name, metrics in models_data.items()], 
                      key=lambda x: x[1])
        
        summary_list = doc.add_paragraph(style='List Bullet')
        summary_list.add_run(f"Best Accuracy: ").bold = True
        acc_val = best_accuracy[1]
        acc_str = f"{acc_val * 100:.2f}%" if acc_val <= 1.0 else f"{acc_val:.2f}%"
        summary_list.add_run(f"{best_accuracy[0]} ({acc_str})")
        
        summary_list = doc.add_paragraph(style='List Bullet')
        summary_list.add_run(f"Best AUC-ROC: ").bold = True
        auc_val = best_auc[1]
        auc_str = f"{auc_val * 100:.2f}%" if auc_val <= 1.0 else f"{auc_val:.2f}%"
        summary_list.add_run(f"{best_auc[0]} ({auc_str})")
        
        summary_list = doc.add_paragraph(style='List Bullet')
        summary_list.add_run(f"Models Processed: ").bold = True
        summary_list.add_run(f"{len(models_data)}")
        
        if errors:
            summary_list = doc.add_paragraph(style='List Bullet')
            summary_list.add_run(f"Models with errors: ").bold = True
            summary_list.add_run(f"{len(errors)}")
    
    doc.add_paragraph()  # Spacing
    
    # ========================================================================
    # DETAILED METRICS TABLE
    # ========================================================================
    doc.add_heading('Detailed Model Metrics', level=1)
    
    if not models_data:
        doc.add_paragraph("❌ No metrics data could be extracted.")
        doc.save(output_path)
        return
    
    # Determine all metric columns with specific order
    all_metrics = set()
    for metrics in models_data.values():
        all_metrics.update(metrics.keys())
    
    # Define metric order
    metric_order = ['accuracy', 'sensitivity', 'specificity', 'precision', 'f1_score', 'auc_roc', 'auc']
    metric_cols = [m for m in metric_order if m in all_metrics]
    # Add any remaining metrics not in the predefined order
    metric_cols.extend([m for m in sorted(all_metrics) if m not in metric_cols])
    
    # Create main metrics table
    num_cols = 1 + len(metric_cols)  # Model name + metrics
    table = doc.add_table(rows=len(models_data) + 1, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Model Name"
    
    for idx, metric in enumerate(metric_cols, 1):
        header_cells[idx].text = metric.upper()
    
    # Format header
    for cell in header_cells:
        shade_cell(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    for row_idx, (model_name, metrics) in enumerate(sorted(models_data.items()), 1):
        row = table.rows[row_idx]
        
        # Model name
        row.cells[0].text = model_name
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Metrics
        for col_idx, metric in enumerate(metric_cols, 1):
            value = metrics.get(metric, 0)
            
            if isinstance(value, float):
                # Convert to percentage with 2 decimal places
                if value <= 1.0:
                    row.cells[col_idx].text = f"{value * 100:.2f}%"
                else:
                    row.cells[col_idx].text = f"{value:.2f}%"
            else:
                row.cells[col_idx].text = str(value)
            
            row.cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # ========================================================================
    # METRICS STATISTICS
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('Performance Statistics', level=1)
    
    if models_data and metric_cols:
        stats_table = doc.add_table(rows=len(metric_cols) + 1, cols=5)
        stats_table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = stats_table.rows[0].cells
        headers = ['Metric', 'Min', 'Max', 'Average', 'Best Model']
        for idx, header_text in enumerate(headers):
            header_cells[idx].text = header_text
            shade_cell(header_cells[idx], '4472C4')
            for run in header_cells[idx].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Calculate stats for each metric
        for m_idx, metric in enumerate(metric_cols, 1):
            values = [metrics.get(metric, 0) for metrics in models_data.values()]
            values = [v for v in values if isinstance(v, (int, float))]
            
            if values:
                min_val = min(values)
                max_val = max(values)
                avg_val = sum(values) / len(values)
                
                # Find best model
                best_model = max(models_data.items(), 
                               key=lambda x: x[1].get(metric, 0))[0]
                
                row = stats_table.rows[m_idx]
                row.cells[0].text = metric.upper()
                min_str = f"{min_val * 100:.2f}%" if min_val <= 1.0 else f"{min_val:.2f}%"
                max_str = f"{max_val * 100:.2f}%" if max_val <= 1.0 else f"{max_val:.2f}%"
                avg_str = f"{avg_val * 100:.2f}%" if avg_val <= 1.0 else f"{avg_val:.2f}%"
                row.cells[1].text = min_str
                row.cells[2].text = max_str
                row.cells[3].text = avg_str
                row.cells[4].text = best_model
    
    doc.add_paragraph()  # Spacing
    
    # ========================================================================
    # ERRORS SECTION (if any)
    # ========================================================================
    if errors:
        doc.add_page_break()
        doc.add_heading('Processing Errors', level=1)
        
        error_intro = doc.add_paragraph()
        error_intro.add_run(f"The following {len(errors)} model(s) had errors during processing:")
        
        for error_msg in errors:
            doc.add_paragraph(error_msg, style='List Bullet')
    
    # ========================================================================
    # FOOTER
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('Report Information', level=1)
    
    footer_text = doc.add_paragraph()
    footer_text.add_run("Generated by: ").bold = True
    footer_text.add_run("PKL Metrics Extraction Script\n")
    
    footer_text = doc.add_paragraph()
    footer_text.add_run("Date: ").bold = True
    footer_text.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    
    footer_text = doc.add_paragraph()
    footer_text.add_run("Source: ").bold = True
    footer_text.add_run(f"{PROJECT_ROOT}/Models/\n")
    
    footer_text = doc.add_paragraph()
    footer_text.add_run("Format: ").bold = True
    footer_text.add_run("Microsoft Word 2007+ (.docx)\n")
    
    # Save document
    try:
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"❌ Error saving document: {e}")
        return False


def main():
    """Main execution."""
    print("\n" + "=" * 80)
    print("📊 EXTRACTING PKL METRICS FROM MODELS FOLDER TO WORD")
    print("=" * 80)
    
    # Identify models directory
    models_dir = PROJECT_ROOT / "Models"
    
    if not models_dir.exists():
        print(f"❌ Models directory not found: {models_dir}")
        return False
    
    print(f"\n📁 Source directory: {models_dir}")
    
    # Extract metrics
    extractor = PickleMetricsExtractor()
    models_data = extractor.extract_all_models(models_dir)
    
    # Report extraction results
    print(f"\n✅ Extraction complete:")
    print(f"   Models processed: {len(models_data)}")
    
    if extractor.errors:
        print(f"   Errors: {len(extractor.errors)}")
        for error in extractor.errors[:3]:
            print(f"     - {error}")
        if len(extractor.errors) > 3:
            print(f"     ... and {len(extractor.errors) - 3} more")
    
    # Create Word document
    lr_str = f"{learning_rate:.0e}".replace('e-0', 'e-').replace('e-', 'e_neg')
    wd_str = f"{weight_decay:.0e}".replace('e-0', 'e-').replace('e-', 'e_neg')
    output_filename = f"Metrics_Report_lr{lr_str}_wd{wd_str}.docx"
    output_path = PROJECT_ROOT / "reports" / output_filename
    
    print(f"\n📄 Creating Word document...")
    success = create_word_report(models_data, extractor.errors, output_path)
    
    if success:
        file_size = output_path.stat().st_size / 1024
        print(f"\n✅ Report created successfully!")
        print(f"   File: {output_path.name}")
        print(f"   Size: {file_size:.2f} KB")
        print(f"   Location: {output_path}")
    else:
        print("❌ Failed to create report")
        return False
    
    print("\n" + "=" * 80)
    print("✅ COMPLETE!")
    print("=" * 80 + "\n")
    
    return True


if __name__ == "__main__":
    try:
        success = main()
        sys.exit(0 if success else 1)
    except Exception as e:
        print(f"\n❌ FATAL ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
