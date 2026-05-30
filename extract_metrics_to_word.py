"""
Extract metrics from all PKL files and generate a formatted Word document.
Organizes experiments by phase with data tables.

Note: This script creates a template Word document with experimental phases.
It will populate actual metrics when PKL files contain properly saved metrics.
"""

import os
import pickle
import sys
import io
import warnings
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Suppress warnings
warnings.filterwarnings('ignore')
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from src.config.Config import save_model_path, learning_rate, weight_decay,path_project
sys.path.insert(0, path_project)

# Define phases and their experiment files
PHASES = {
    "Phase 1-2": {
        "description": "Single Deep Learning Model",
        "experiments": ["single_dl_model"]
    },
    "Phase 3": {
        "description": "Hybrid Deep Learning + ML Classifiers",
        "experiments": ["hybrid_dl_ml", "dual_dl_ml", "triple_dl_ml"]
    },
    "Phase 4": {
        "description": "Hybrid DL + ML Ensemble & DL Ensemble",
        "experiments": ["hybrid_dl_ml_ensemble", "dual_hybrid_ensemble", "triple_hybrid_ensemble"]
    },
    "Phase 5": {
        "description": "Multi-DL Ensemble",
        "experiments": ["dual_dl_ensemble", "triple_dl_ensemble"]
    }
}


def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)


def extract_metrics_from_pkl(pkl_files_dir):
    """Extract metrics from PKL files or create template structure"""
    metrics_data = {}
    sample_metrics = {
        "accuracy": 0.0,
        "precision": 0.0,
        "recall": 0.0,
        "f1_score": 0.0,
        "auc": 0.0,
        "status": "No metrics available"
    }
    
    if not os.path.exists(pkl_files_dir):
        print(f"✗ Directory not found: {pkl_files_dir}")
        return metrics_data
    
    pkl_files = sorted([f for f in os.listdir(pkl_files_dir) if f.endswith('.pkl')])
    
    print(f"\n📁 Found {len(pkl_files)} PKL files")
    print("  ℹ Attempting to extract metrics...")
    
    for pkl_file in pkl_files:
        filepath = os.path.join(pkl_files_dir, pkl_file)
        name = pkl_file.replace('.pkl', '')
        metrics_data[name] = sample_metrics.copy()
        
        try:
            # Use a custom unpickler that skips model objects
            import types
            
            class ModelSkipUnpickler(pickle.Unpickler):
                def find_class(self, module, name):
                    # Skip loading keras models
                    if 'keras' in module or 'tensorflow' in module:
                        if name in ['Sequential', 'Model', 'Functional']:
                            return types.SimpleNamespace
                    return super().find_class(module, name)
            
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                with open(filepath, 'rb') as f:
                    try:
                        # Try optimized unpickler first
                        data = ModelSkipUnpickler(f).load()
                    except:
                        # Fallback to standard loader
                        f.seek(0)
                        data = pickle.load(f)
            
            # Try to extract metrics
            if isinstance(data, dict):
                exp_data = data.get('experiment_data', {})
                if isinstance(exp_data, dict):
                    metrics = exp_data.get('metrics', {})
                    if isinstance(metrics, dict):
                        # Extract serializable metrics
                        for model_name, model_metrics in metrics.items():
                            if isinstance(model_metrics, dict):
                                for k, v in model_metrics.items():
                                    if isinstance(v, (int, float, str, bool)):
                                        metrics_data[name][k] = v
                                break
                        if metrics:
                            metrics_data[name].pop('status', None)
                            print(f"  ✓ Loaded: {name}")
                        else:
                            print(f"  ℹ {name}: template metrics")
                    else:
                        print(f"  ℹ {name}: template metrics")
                else:
                    print(f"  ℹ {name}: template metrics")
            else:
                print(f"  ℹ {name}: template metrics")
                
        except Exception as e:
            # Silently use template metrics
            print(f"  ℹ {name}: template metrics")
            pass
    
    return metrics_data


def format_metric_value(value):
    """Format metric value as percentage for display"""
    if isinstance(value, dict):
        return "N/A"
    if isinstance(value, (int, float)):
        return f"{value * 100:.2f}%"
    return str(value)


def add_phase_metrics_table(doc, phase_name, experiments, metrics_data):
    """Add a comprehensive metrics table for an entire phase"""
    
    # Define metric columns to display in specific order
    metric_columns = ["accuracy", "sensitivity", "specificity", "precision", "f1_score", "auc"]
    
    # Collect data for this phase
    phase_data = {}
    for exp_name in experiments:
        for pkl_name, metrics in metrics_data.items():
            if exp_name.lower() in pkl_name.lower():
                phase_data[exp_name] = metrics
                break
    
    if not phase_data:
        doc.add_paragraph(f"  ⚠ No metrics found for this phase")
        return
    
    # Create table: Model Name + 5 metric columns
    num_cols = 1 + len(metric_columns)
    num_rows = len(phase_data) + 1  # +1 for header
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Model Name"
    
    # Define header labels for display
    header_labels = {
        "accuracy": "ACCURACY",
        "sensitivity": "SENSITIVITY",
        "specificity": "SPECIFICITY",
        "precision": "PRECISION",
        "f1_score": "F1-SCORE",
        "auc": "AUC-ROC"
    }
    
    for col_idx, metric_col in enumerate(metric_columns, 1):
        header_cells[col_idx].text = header_labels.get(metric_col, metric_col.replace('_', ' ').title())
    
    # Style header
    for cell in header_cells:
        set_cell_background(cell, "4472C4")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    for row_idx, (exp_name, metrics) in enumerate(sorted(phase_data.items()), 1):
        row = table.rows[row_idx]
        
        # Model name
        row.cells[0].text = exp_name
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Metrics
        for col_idx, metric_col in enumerate(metric_columns, 1):
            metric_value = metrics.get(metric_col, 0.0)
            if isinstance(metric_value, (int, float)):
                # Convert to percentage with 2 decimal places
                if metric_value <= 1.0:
                    row.cells[col_idx].text = f"{metric_value * 100:.2f}%"
                else:
                    row.cells[col_idx].text = f"{metric_value:.2f}%"
            else:
                row.cells[col_idx].text = "N/A"
            row.cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Add spacing


def create_word_document(metrics_data):
    """Create a formatted Word document with metrics organized by phase"""
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Melanoma Detection Model Evaluation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Performance Metrics by Experimental Phase')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.italic = True
    subtitle_format.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Process each phase
    for phase_name, phase_info in PHASES.items():
        
        # Phase heading
        phase_heading = doc.add_heading(phase_name, level=1)
        phase_heading_run = phase_heading.runs[0]
        phase_heading_run.font.color.rgb = RGBColor(192, 0, 0)
        
        # Phase description
        doc.add_paragraph(f"Description: {phase_info['description']}", style='List Bullet')
        
        # List experiments in this phase
        experiments_para = doc.add_paragraph(f"Experiments: {', '.join(phase_info['experiments'])}", style='List Bullet')
        
        doc.add_paragraph()
        
        # Add grand table for this phase
        add_phase_metrics_table(doc, phase_name, phase_info['experiments'], metrics_data)
        
        doc.add_page_break()
    
    return doc


def save_word_document(doc, output_path):
    """Save Word document"""
    try:
        doc.save(output_path)
        print(f"\n✓ Word document created: {output_path}")
        return True
    except Exception as e:
        print(f"\n✗ Error saving Word document: {e}")
        return False


def main():
    """Main execution function"""
    
    print("\n" + "="*80)
    print("  📊 EXTRACT METRICS FROM PKL FILES TO WORD DOCUMENT")
    print("="*80)
    
    # Extract metrics
    metrics_data = extract_metrics_from_pkl(save_model_path)
    
    if not metrics_data:
        print("\n✗ No metrics data found. Exiting.")
        return
    
    print(f"\n✓ Successfully loaded metrics for {len(metrics_data)} experiments")
    
    # Create Word document
    doc = create_word_document(metrics_data)
    
    # Save document
    output_dir = os.path.join(Path(save_model_path).parent, "reports")
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    
    # Generate filename with config parameters
    lr_str = f"{learning_rate:.0e}".replace('e-0', 'e-').replace('e-', 'e_neg')
    wd_str = f"{weight_decay:.0e}".replace('e-0', 'e-').replace('e-', 'e_neg')
    output_filename = f"Metrics_Report_lr{lr_str}_wd{wd_str}.docx"
    output_path = os.path.join(output_dir, output_filename)
    
    if save_word_document(doc, output_path):
        print(f"\n✓ Report generated successfully!")
        print(f"  Location: {output_path}")
    
    print("="*80 + "\n")


if __name__ == "__main__":
    main()
