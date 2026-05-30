"""
Extract images and metrics from PKL files and generate a formatted Word document.
Organizes all experiments with visualizations by filename.
Supports multiple models inside a single PKL file.
"""

import os
import pickle
import sys

import warnings
import matplotlib.pyplot as plt
import numpy as np
import io
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sklearn.metrics import confusion_matrix
import seaborn as sns
from src.modules.memory_optimized_test import test_memory_optimized

test_memory_optimized()

warnings.filterwarnings('ignore')
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
try:
    from src.config.Config import learning_rate, weight_decay,path_project
    sys.path.insert(0, path_project)
except ImportError:
    learning_rate = 3e-4
    weight_decay = 1e-5

plt.ioff()


def set_cell_background(cell, fill):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)


# ------------------------------------------------------------
# PKL EXTRACTION (MULTI MODEL)
# ------------------------------------------------------------

def extract_pkl_data(pkl_file):

    try:

        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            with open(pkl_file, 'rb') as f:
                data = pickle.load(f)

        exp_data = data.get('experiment_data', {})
        history = exp_data.get('history', {})

        if hasattr(history, "history"):
            history = history.history

        models = []

        metrics = exp_data.get("metrics", {})

        if isinstance(metrics, dict):

            for model_name, value in metrics.items():

                result = {
                    'name': model_name,
                    'metrics': {},
                    'history': history,
                    'y_true': None,
                    'y_pred': None,
                    'success': True
                }

                if isinstance(value, dict):

                    if "y_true" in value:
                        result["y_true"] = value["y_true"]

                    if "y_pred" in value:
                        result["y_pred"] = value["y_pred"]

                    for k, v in value.items():

                        if k in ["y_true", "y_pred"]:
                            continue

                        if isinstance(v, (int, float, np.integer, np.floating)):
                            result["metrics"][k] = float(v)

                        elif isinstance(v, (str, bool)):
                            result["metrics"][k] = v

                models.append(result)

        return {
            "models": models,
            "success": True
        }

    except Exception as e:

        print(f"⚠ Error loading {pkl_file}: {e}")

        return {
            "success": False,
            "error": str(e)
        }


# ------------------------------------------------------------
# PLOTS
# ------------------------------------------------------------

def create_training_history_plot(history, filename):

    if not history:
        return None

    try:

        fig, axes = plt.subplots(1, 2, figsize=(12, 4))
        fig.suptitle(f'Training History - {filename}', fontsize=14, fontweight='bold')

        
        hist_data = {}

        for key, values in history.items():

            if hasattr(values, "tolist"):
                hist_data[key] = values.tolist()
            else:
                hist_data[key] = values
        # Plot 1: Loss
        has_loss = False
        if "loss" in hist_data:

            epochs = range(1, len(hist_data["loss"]) + 1)

            axes[0].plot(epochs, hist_data['loss'], 'b-', label='Train Loss', linewidth=2)
            has_loss = True
            if "val_loss" in hist_data:
                axes[0].plot(epochs, hist_data["val_loss"], 'r-', label="Val Loss", linewidth=2)
        if has_loss:
            axes[0].set_xlabel('Epoch')
            axes[0].set_ylabel('Loss')
            axes[0].set_title("Loss")
            axes[0].legend()
            axes[0].grid(True, alpha=0.3)

        # Plot 2: Accuracy
        has_acc = False
        if "accuracy" in hist_data:

            epochs = range(1, len(hist_data["accuracy"]) + 1)

            axes[1].plot(epochs, hist_data['accuracy'], 'b-', label='Train Accuracy', linewidth=2)
            has_acc = True
            if "val_accuracy" in hist_data:
                axes[1].plot(epochs, hist_data["val_accuracy"], 'r-', label="Val Accuracy", linewidth=2)
        if has_acc:
            axes[1].set_xlabel('Epoch')
            axes[1].set_ylabel('Accuracy')
            axes[1].set_title("Accuracy")
            axes[1].legend()
            axes[1].grid(True, alpha=0.3)

        if not (has_loss or has_acc):
            return None
        plt.tight_layout()

        # Save to bytes
        img = io.BytesIO()

        plt.savefig(img, format="png", dpi=100, bbox_inches="tight")

        img.seek(0)

        plt.close(fig)

        return img

    except Exception as e:

        print(f"    Error creating plot: {e}")

        return None


def create_confusion_matrix_plot(y_true, y_pred, filename):
    """Create confusion matrix plot with counts and percentages"""
    if y_true is None or y_pred is None:
        return None
    
    try:
        cm = confusion_matrix(y_true, y_pred)
        
        # Calculate total and percentages
        total = cm.sum()
        percentages = cm.astype('float') / total * 100
        
        # Create annotations with both counts and percentages
        annotations = np.empty_like(cm, dtype=object)
        for i in range(cm.shape[0]):
            for j in range(cm.shape[1]):
                annotations[i, j] = f'{cm[i, j]}\n({percentages[i, j]:.2f}%)'
        
        fig, ax = plt.subplots(figsize=(8, 6))
        sns.heatmap(cm, annot=annotations, fmt='', cmap='Blues', cbar=True, ax=ax,
                    xticklabels=['Benign', 'Malignant'],
                    yticklabels=['Benign', 'Malignant'],
                    cbar_kws={'label': 'Count'})
        
        ax.set_title(f'Confusion Matrix - {filename}', fontsize=14, fontweight='bold')
        ax.set_xlabel('Predicted Label')
        ax.set_ylabel('True Label')
        
        plt.tight_layout()
        
        # Save to bytes
        img_bytes = io.BytesIO()
        plt.savefig(img_bytes, format='png', dpi=100, bbox_inches='tight')
        img_bytes.seek(0)
        plt.close(fig)
        
        return img_bytes
    except Exception as e:
        print(f"    Error creating confusion matrix: {e}")
        return None


def create_metrics_bar_plot(metrics, filename):

    if not metrics:
        return None

    try:

        keys = ['accuracy','precision','sensitivity','specificity','f1-score','auc-roc']

        data = {k:metrics[k] for k in keys if k in metrics}

        if not data:
            return None

        names = list(data.keys())
        values = list(data.values())

        fig, ax = plt.subplots(figsize=(8,4))

        # Colors for each bar
        colors = [
            "#1f77b4",  # blue
            "#ff7f0e",  # orange
            "#2ca02c",  # green
            "#d62728",  # red
            "#9467bd",  # purple
            "#8c564b",  # brown
            "#e377c2",  # pink
            "#7f7f7f"   # gray
        ]

        bar_colors = colors[:len(values)]

        bars = ax.bar(names, values, color=bar_colors)

        # add value on top of bars
        for bar in bars:

            height = bar.get_height()

            ax.text(
                bar.get_x() + bar.get_width()/2,
                height,
                f"{height:.2f}%",
                ha='center',
                va='bottom',
                fontsize=10
            )

        ax.set_ylim(0, max(values)*1.2)

        ax.set_title(f"Metrics - {filename}")

        ax.set_ylabel("Score (%)")

        ax.grid(axis='y', alpha=0.3)

        plt.xticks(rotation=45)

        plt.tight_layout()

        img = io.BytesIO()

        plt.savefig(img, format="png", dpi=100, bbox_inches="tight")

        img.seek(0)

        plt.close()

        return img

    except Exception as e:

        print("Metrics plot error:", e)

        return None

# ------------------------------------------------------------
# WORD SECTION
# ------------------------------------------------------------

def add_experiment_section(doc, filename, pkl_data):
    """Add experiment section to document"""
    
    # Experiment heading
    exp_heading = doc.add_heading(filename, level=2)
    exp_heading_run = exp_heading.runs[0]
    exp_heading_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Model name
    model_name = pkl_data.get('name', 'Unknown')
    doc.add_paragraph(f"Model: {model_name}", style='List Bullet')
    
    # Metrics table
    metrics = pkl_data.get('metrics', {})
    if metrics:
        doc.add_paragraph("Performance Metrics:", style='Heading 3')
        
        table = doc.add_table(rows=len(metrics) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = "Metric"
        header_cells[1].text = "Value"
        
        for cell in header_cells:
            set_cell_background(cell, "4472C4")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
        
        # Data rows
        for row_idx, (metric_name, metric_value) in enumerate(sorted(metrics.items()), 1):
            row = table.rows[row_idx]
            row.cells[0].text = metric_name.replace('_', ' ').replace('-', ' ').title()
            
            if isinstance(metric_value, float):
                # Always display as percentage with 2 decimal places
                row.cells[1].text = f"{metric_value:.2f}%"
            else:
                row.cells[1].text = str(metric_value)
        
        doc.add_paragraph()
    
    # Metrics Bar Plot
    metrics_plot = create_metrics_bar_plot(metrics, filename)
    if metrics_plot:
        doc.add_paragraph("Metrics Comparison:", style='Heading 3')
        try:
            doc.add_picture(metrics_plot, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"    Error adding metrics plot: {e}")
        doc.add_paragraph()
    
    # Confusion Matrix
    y_true = pkl_data.get('y_true')
    y_pred = pkl_data.get('y_pred')
    if y_true is not None and y_pred is not None:
        cm_plot = create_confusion_matrix_plot(y_true, y_pred, filename)
        if cm_plot:
            doc.add_paragraph("Confusion Matrix:", style='Heading 3')
            try:
                doc.add_picture(cm_plot, width=Inches(4.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"    Error adding confusion matrix: {e}")
            doc.add_paragraph()
    
    # Training history plot
    history = pkl_data.get('history', {})
    if history:
        doc.add_paragraph("Training History:", style='Heading 3')
        
        plot_img = create_training_history_plot(history, filename)
        if plot_img:
            try:
                doc.add_picture(plot_img, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"    Error adding plot: {e}")
        
        doc.add_paragraph()
    
    # PAGE BREAK AFTER EACH MODEL
    doc.add_page_break()

# ------------------------------------------------------------
# DOCUMENT CREATION
# ------------------------------------------------------------

def create_word_document(pkl_files_dir, output_path):

    if not os.path.exists(pkl_files_dir):

        print("Directory not found:", pkl_files_dir)

        return False

    pkl_files = sorted(
        [f for f in os.listdir(pkl_files_dir) if f.endswith(".pkl")]
    )

    if not pkl_files:

        print("No PKL files found")

        return False

    doc = Document()

    title = doc.add_heading("PKL Models Extraction Report",0)

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    for pkl_file in pkl_files:

        filename = pkl_file.replace(".pkl","")

        filepath = os.path.join(pkl_files_dir,pkl_file)

        print("Processing:", filename)

        data = extract_pkl_data(filepath)

        if data.get("success"):

            models = data.get("models",[])

            for model in models:

                section_name = f"{filename} - {model['name']}"

                add_experiment_section(doc, section_name, model)

        else:

            doc.add_heading(filename,2)

            doc.add_paragraph(data.get("error"))

    doc.save(output_path)

    print("Report saved:", output_path)

    return True


# ------------------------------------------------------------
# MAIN
# ------------------------------------------------------------

def main():

     """Main execution function"""
    
     print("\n" + "="*80)
     print("  📊 EXTRACT PKL DATA TO WORD DOCUMENT")
     print("="*80)
    
    # Paths
     pkl_dir = "/home/akachat/tf_env/Ensemble_CNN_Melanoma/Models"
     output_dir = "/home/akachat/tf_env/Ensemble_CNN_Melanoma/reports"
    
    # Create output directory
     Path(output_dir).mkdir(parents=True, exist_ok=True)
    
    # Generate filename with config parameters
     lr_str = f"{learning_rate:.0e}".replace('e-0', 'e-').replace('e-', 'e_neg')
     wd_str = f"{weight_decay:.0e}".replace('e-0', 'e-').replace('e-', 'e_neg')
     output_filename = f"PKL_Models_Report_lr{lr_str}_wd{wd_str}.docx"
     output_path = os.path.join(output_dir, output_filename)
    
    # Create document
     success = create_word_document(pkl_dir, output_path)
    
     if success:
        print(f"\n✓ Report generated successfully!")
        print(f"  Location: {output_path}")
    
     print("="*80 + "\n")


if __name__ == "__main__":
    main()